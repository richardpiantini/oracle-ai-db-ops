import json
import re
import subprocess
import time
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import oci
import oracledb
from docx import Document
from fastapi import FastAPI
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook
from pydantic import BaseModel

from config import *

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"
TOOLS_FILE = BASE_DIR / "tools.json"
EXPORT_DIR = BASE_DIR / "exports"
EXPORT_DIR.mkdir(exist_ok=True)
DATABASES_FILE = BASE_DIR / "databases.json"

# Connection pool
pool = oracledb.create_pool(
    user=DB_USER, password=DB_PASS, dsn=DB_DSN,
    min=2, max=5, increment=1
)

app = FastAPI(title="Oracle AI Database Operations")

# Multi-database pool manager
_db_pools: dict[str, oracledb.ConnectionPool] = {}

def load_databases():
    return read_json(DATABASES_FILE, [])

def get_pool(db_id: str | None = None) -> oracledb.ConnectionPool:
    if not db_id:
        return pool  # default pool

    if db_id in _db_pools:
        return _db_pools[db_id]

    databases = load_databases()
    db = next((d for d in databases if d["id"] == db_id), None)
    if not db:
        raise ValueError(f"Database not found: {db_id}")
    if not db.get("active"):
        raise ValueError(f"Database not available: {db['name']}. Contact your admin to configure this connection.")

    dsn = f"{db['host']}:{db['port']}/{db['service']}"
    new_pool = oracledb.create_pool(
        user=db["user"], password=db["password"], dsn=dsn,
        min=2, max=5, increment=1
    )
    _db_pools[db_id] = new_pool
    return new_pool

RUN_HISTORY: dict[str, dict[str, Any]] = {}

# -----------------------------
# Utilities
# -----------------------------

def read_json(path: Path, default: Any):
    if not path.exists():
        return default
    return json.loads(path.read_text())

def write_json(path: Path, data: Any):
    path.write_text(json.dumps(data, indent=2))

def load_tools():
    return read_json(TOOLS_FILE, [])

def save_tools(tools):
    write_json(TOOLS_FILE, tools)

def next_tool_id(tools):
    nums = [int(t["id"]) for t in tools if str(t.get("id", "")).isdigit()]
    return str(max(nums) + 1 if nums else 1)

def is_safe_select_sql(sql: str) -> bool:
    sql_clean = (sql or "").strip().rstrip(";").strip()
    lowered = sql_clean.lower()

    if not lowered.startswith("select") and not lowered.startswith("with"):
        return False

    blocked = [
        " insert ", " update ", " delete ", " merge ", " alter ", " drop ",
        " truncate ", " create ", " grant ", " revoke ", " commit ",
        " rollback ", " execute ", " begin ", " declare ", " call "
    ]
    padded = f" {lowered} "
    return not any(tok in padded for tok in blocked)

def query_oracle(sql, params=None, db_id=None):
    p = get_pool(db_id)
    with p.acquire() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, params or {})
            cols = [d[0] for d in cur.description]
            return [dict(zip(cols, row)) for row in cur.fetchall()]

def parse_dt(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value)
    except Exception:
        return None

def fmt_ts(dt: datetime) -> str:
    return dt.strftime("%Y-%m-%d %H:%M:%S")

def build_time_window(start_time: str | None, end_time: str | None, hours_back: int | None):
    now = datetime.now()

    start_dt = parse_dt(start_time)
    end_dt = parse_dt(end_time)

    if start_dt and end_dt and start_dt < end_dt:
        return {
            "start_dt": start_dt,
            "end_dt": end_dt,
            "start_str": fmt_ts(start_dt),
            "end_str": fmt_ts(end_dt),
            "label": f"{fmt_ts(start_dt)} to {fmt_ts(end_dt)}"
        }

    if hours_back and hours_back > 0:
        start_dt = now - timedelta(hours=hours_back)
        end_dt = now
        return {
            "start_dt": start_dt,
            "end_dt": end_dt,
            "start_str": fmt_ts(start_dt),
            "end_str": fmt_ts(end_dt),
            "label": f"last {hours_back} hours"
        }

    return None

def apply_time_window_sql(sql: str, window: dict[str, Any] | None) -> tuple[str, dict]:
    if not window:
        return sql, {}

    params = {
        "start_ts": window["start_str"],
        "end_ts": window["end_str"],
    }

    replacements = [
        (
            r"s\.end_interval_time\s*>=\s*SYSTIMESTAMP\s*-\s*INTERVAL\s*'[^']+'\s*HOUR",
            "s.end_interval_time BETWEEN TO_TIMESTAMP(:start_ts, 'YYYY-MM-DD HH24:MI:SS') AND TO_TIMESTAMP(:end_ts, 'YYYY-MM-DD HH24:MI:SS')",
        ),
        (
            r"sample_time\s*>=\s*SYSTIMESTAMP\s*-\s*INTERVAL\s*'[^']+'\s*HOUR",
            "sample_time BETWEEN TO_TIMESTAMP(:start_ts, 'YYYY-MM-DD HH24:MI:SS') AND TO_TIMESTAMP(:end_ts, 'YYYY-MM-DD HH24:MI:SS')",
        ),
        (
            r"begin_time\s*>=\s*SYSDATE\s*-\s*1/24",
            "begin_time BETWEEN TO_DATE(:start_ts, 'YYYY-MM-DD HH24:MI:SS') AND TO_DATE(:end_ts, 'YYYY-MM-DD HH24:MI:SS')",
        ),
        (
            r"timestamp\s*>=\s*SYSDATE\s*-\s*1",
            "timestamp BETWEEN TO_DATE(:start_ts, 'YYYY-MM-DD HH24:MI:SS') AND TO_DATE(:end_ts, 'YYYY-MM-DD HH24:MI:SS')",
        ),
    ]

    out = sql
    changed = False
    for pattern, repl in replacements:
        new_sql = re.sub(pattern, repl, out, flags=re.IGNORECASE)
        if new_sql != out:
            changed = True
            out = new_sql

    if not changed:
        return sql, {}

    used_params = {}
    if ":start_ts" in out:
        used_params["start_ts"] = window["start_str"]
    if ":end_ts" in out:
        used_params["end_ts"] = window["end_str"]

    return out, used_params

def extract_views(sql: str) -> list[str]:
    matches = re.findall(r"\b(?:from|join)\s+([a-zA-Z0-9_\$#]+)", sql, flags=re.IGNORECASE)
    return list(dict.fromkeys([m.upper() for m in matches]))

ALLOWED_VIEWS = {
    "V$SESSION", "V$SQL", "V$SQL_PLAN", "V$LOCK", "V$SYSSTAT", "V$SYSTEM_EVENT",
    "V$PROCESS", "V$RESOURCE_LIMIT", "V$PGASTAT", "V$SQL_WORKAREA_ACTIVE",
    "V$UNDOSTAT", "V$TEMP_SPACE_HEADER", "V$DATABASE", "V$DATAGUARD_STATUS",
    "V$ARCHIVE_DEST_STATUS", "V$MANAGED_STANDBY", "V$DATAGUARD_STATS",
    "V$SESSION_WAIT", "V$SQLAREA", "V$RMAN_BACKUP_JOB_DETAILS", "V$BACKUP_SET",
    "V$BACKUP_PIECE", "V$ARCHIVED_LOG", "V$ARCHIVE_DEST", "DBA_HIST_ACTIVE_SESS_HISTORY",
    "DBA_HIST_SYSTEM_EVENT", "DBA_HIST_SYSSTAT", "DBA_HIST_SNAPSHOT",
    "DBA_HIST_SQLSTAT", "DBA_HIST_SQL_PLAN", "DBA_OBJECTS", "DBA_SEGMENTS",
    "DBA_TABLESPACE_USAGE_METRICS", "DBA_TEMP_FREE_SPACE", "DBA_TABLESPACES",
    "DBA_HIST_TBSPC_SPACE_USAGE", "DUAL"
}

REFERENCE_HINTS = {
    "awr": ["skills/performance/awr-reports.md"],
    "ash": ["skills/performance/ash-analysis.md"],
    "wait": ["skills/performance/wait-events.md"],
    "sql plan": ["skills/performance/explain-plan.md"],
    "execution plan": ["skills/performance/explain-plan.md"],
    "optimizer": ["skills/performance/optimizer-stats.md"],
    "stats": ["skills/performance/optimizer-stats.md"],
    "index": ["skills/performance/index-strategy.md"],
    "lock": ["skills/appdev/locking-concurrency.md", "skills/performance/wait-events.md"],
    "blocking": ["skills/appdev/locking-concurrency.md", "skills/performance/wait-events.md"],
    "concurrency": ["skills/appdev/locking-concurrency.md", "skills/performance/wait-events.md"],
    "memory": ["skills/performance/memory-tuning.md"],
    "pga": ["skills/performance/memory-tuning.md"],
    "undo": ["skills/admin/undo-management.md"],
    "tablespace": ["skills/admin/undo-management.md", "skills/admin/backup-recovery.md"],
    "backup": ["skills/admin/backup-recovery.md"],
    "recovery": ["skills/admin/backup-recovery.md"],
    "rman": ["skills/admin/backup-recovery.md"],
    "dataguard": ["skills/admin/dataguard.md", "skills/admin/redo-log-management.md"],
    "data guard": ["skills/admin/dataguard.md", "skills/admin/redo-log-management.md"],
    "redo": ["skills/admin/redo-log-management.md"],
    "transport lag": ["skills/admin/dataguard.md", "skills/admin/redo-log-management.md"],
    "standby": ["skills/admin/dataguard.md"],
}

def infer_references(name: str, prompt: str, sql_mode: str, sql: str = "", queries: dict | None = None):
    text = " ".join([
        name or "",
        prompt or "",
        sql or "",
        " ".join((queries or {}).keys()),
        " ".join((queries or {}).values()),
    ]).lower()

    found = []
    for key, refs in REFERENCE_HINTS.items():
        if key in text:
            for ref in refs:
                if ref not in found:
                    found.append(ref)
    return found

def validate_tool_definition(tool: dict[str, Any]) -> tuple[bool, str]:
    sql_mode = tool.get("sql_mode")
    if sql_mode not in ["single", "multi"]:
        return False, "Unsupported tool mode"

    sqls = []
    if sql_mode == "single":
        sql = (tool.get("sql") or "").strip()
        if not sql or not is_safe_select_sql(sql):
            return False, "Single-query tool must contain a safe SELECT/CTE query"
        sqls.append(sql)
    else:
        queries = tool.get("queries") or {}
        if not queries:
            return False, "Multi-query tool must contain queries"
        for label, sql in queries.items():
            if not label.strip():
                return False, "Query block label cannot be empty"
            if not is_safe_select_sql(sql):
                return False, f"Unsafe SQL in query block: {label}"
            sqls.append(sql)

    for sql in sqls:
        for view in extract_views(sql):
            if view not in ALLOWED_VIEWS:
                return False, f"View not allowed in this environment: {view}"

    return True, "ok"

# -----------------------------
# GenAI
# -----------------------------

SYSTEM_PROMPT = """You are a senior Oracle DBA analyst. You receive structured performance
data from an Oracle database and produce actionable operational analysis.

For every response, use this structure:

SUMMARY
One or two sentences: what is happening and whether it needs attention.

ROOT CAUSE
The primary factor driving the current state. Be specific: name the
wait event, SQL_ID, object, or resource.

EVIDENCE
Cite the specific metrics that support your diagnosis. Use numbers.

SEVERITY
One of: Critical (immediate action needed), Warning (monitor closely),
Informational (normal operations).

TREND
One of: Degrading (getting worse), Stable (holding steady),
Improving (recovering). If insufficient data for trend, say so.

RECOMMENDED ACTIONS
Numbered list of concrete next steps, ordered by impact. Include the
Oracle command or view to check where relevant.

Rules:
- Never fabricate metrics. Only reference data you received.
- If data is insufficient for a conclusion, say so explicitly.
- Prefer specific SQL_IDs, object names, and wait events over generic advice.
- When comparing periods, express changes as percentages and absolutes.
- For predictive analysis, state confidence level and assumptions behind projections.
"""

def ask_genai(prompt, data):
    config = oci.config.from_file()
    client = oci.generative_ai_inference.GenerativeAiInferenceClient(
        config,
        service_endpoint=OCI_GENAI_ENDPOINT
    )

    user_msg = (
        f"{prompt}\n\n"
        f"Here is the live data from Oracle:\n\n"
        f"{json.dumps(data, indent=2, default=str)}"
    )

    response = client.chat(
        oci.generative_ai_inference.models.ChatDetails(
            compartment_id=OCI_COMPARTMENT,
            chat_request=oci.generative_ai_inference.models.GenericChatRequest(
                messages=[
                    oci.generative_ai_inference.models.SystemMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=SYSTEM_PROMPT)]
                    ),
                    oci.generative_ai_inference.models.UserMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=user_msg)]
                    ),
                ],
                max_tokens=3000,
                temperature=0.2,
                is_stream=False,
            ),
            serving_mode=oci.generative_ai_inference.models.OnDemandServingMode(
                model_id=OCI_MODEL_ID
            ),
        )
    )
    return response.data.chat_response.choices[0].message.content[0].text

AI_TOOL_CREATOR_PROMPT = f"""You generate Oracle DBA tool definitions as JSON only.

Return valid JSON with exactly these keys:
- name
- category
- prompt
- sql_mode
- sql
- queries

Rules:
- category must be one of: real-time, predictive, infra
- sql_mode must be one of: single, multi
- if single, sql must contain one safe Oracle SELECT/CTE query and queries must be {{}}
- if multi, queries must be an object of named safe Oracle SELECT/CTE queries and sql must be ""
- do not use INSERT, UPDATE, DELETE, MERGE, ALTER, DROP, TRUNCATE, CREATE, BEGIN, DECLARE, CALL
- only use these Oracle views/tables: {sorted(ALLOWED_VIEWS)}
- for AWR/ASH/performance, prefer DBA_HIST_* views plus V$SYSSTAT / V$SYSTEM_EVENT where appropriate
- for locking, prefer V$SESSION, V$LOCK, DBA_OBJECTS, V$SQL
- for backup/recovery, prefer V$RMAN_BACKUP_JOB_DETAILS, V$BACKUP_SET, V$BACKUP_PIECE, V$ARCHIVED_LOG, V$DATABASE
- for Data Guard, prefer V$DATABASE, V$ARCHIVE_DEST_STATUS, V$DATAGUARD_STATUS, V$DATAGUARD_STATS
- choose the best Oracle skills guidance based on the request, especially from:
  - skills/performance/awr-reports.md
  - skills/performance/ash-analysis.md
  - skills/performance/wait-events.md
  - skills/performance/explain-plan.md
  - skills/performance/index-strategy.md
  - skills/performance/optimizer-stats.md
  - skills/performance/memory-tuning.md
  - skills/appdev/locking-concurrency.md
  - skills/admin/backup-recovery.md
  - skills/admin/dataguard.md
  - skills/admin/redo-log-management.md
  - skills/admin/undo-management.md
- output JSON only with no markdown fences

Key column names (use these exact names, not guesses):
- DBA_HIST_SNAPSHOT: snap_id, dbid, instance_number, begin_interval_time, end_interval_time (NOT begin_time/end_time)
- DBA_HIST_SYSTEM_EVENT: snap_id, dbid, instance_number, event_name (NOT event), wait_class, total_waits, time_waited_micro
- DBA_HIST_SYSSTAT: snap_id, dbid, instance_number, stat_name, value
- DBA_HIST_SQLSTAT: snap_id, dbid, instance_number, sql_id, plan_hash_value, elapsed_time_delta, executions_delta, buffer_gets_delta, disk_reads_delta, rows_processed_delta
- DBA_HIST_ACTIVE_SESS_HISTORY: snap_id, dbid, instance_number, sample_time, session_id, sql_id, event, wait_class, session_state
- V$SESSION: sid, serial#, username, sql_id, event, wait_class, seconds_in_wait, blocking_session, status, type, module, action, row_wait_obj#
- V$SYSSTAT: name, value (NOT stat_name)
- V$SYSTEM_EVENT: event, wait_class, total_waits, time_waited_micro
- V$RESOURCE_LIMIT: resource_name, current_utilization, max_utilization, limit_value
- V$UNDOSTAT: begin_time, end_time, undotsn, undoblks, maxquerylen, unxpstealcnt, ssolderrcnt
- V$SQL: sql_id, sql_text, executions, buffer_gets, disk_reads, rows_processed, elapsed_time, cpu_time, plan_hash_value
- V$SQL_PLAN: sql_id, plan_hash_value, operation, options, object_name, object_owner
- V$RMAN_BACKUP_JOB_DETAILS: session_key, start_time, end_time, status, input_type, elapsed_seconds, output_bytes, output_bytes_display
- DBA_TABLESPACE_USAGE_METRICS: tablespace_name, used_space, tablespace_size, used_percent
- DBA_TEMP_FREE_SPACE: tablespace_name, tablespace_size, free_space
- When joining DBA_HIST_* to DBA_HIST_SNAPSHOT, always use: e.snap_id = s.snap_id AND e.dbid = s.dbid AND e.instance_number = s.instance_number
- Time filters on DBA_HIST_SNAPSHOT use: s.end_interval_time >= SYSTIMESTAMP - INTERVAL 'N' HOUR
- Time filters on V$UNDOSTAT use: begin_time >= SYSDATE - 1/24
"""

def ask_genai_json(user_request: str):
    config = oci.config.from_file()
    client = oci.generative_ai_inference.GenerativeAiInferenceClient(
        config,
        service_endpoint=OCI_GENAI_ENDPOINT
    )

    response = client.chat(
        oci.generative_ai_inference.models.ChatDetails(
            compartment_id=OCI_COMPARTMENT,
            chat_request=oci.generative_ai_inference.models.GenericChatRequest(
                messages=[
                    oci.generative_ai_inference.models.SystemMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=AI_TOOL_CREATOR_PROMPT)]
                    ),
                    oci.generative_ai_inference.models.UserMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=user_request)]
                    ),
                ],
                max_tokens=2500,
                temperature=0.15,
                is_stream=False,
            ),
            serving_mode=oci.generative_ai_inference.models.OnDemandServingMode(
                model_id=OCI_MODEL_ID
            ),
        )
    )

    raw = response.data.chat_response.choices[0].message.content[0].text.strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("AI tool creator did not return valid JSON")
    return json.loads(raw[start:end + 1])

def repair_tool_with_error(user_request: str, broken_tool: dict[str, Any], db_error: str):
    config = oci.config.from_file()
    client = oci.generative_ai_inference.GenerativeAiInferenceClient(
        config,
        service_endpoint=OCI_GENAI_ENDPOINT
    )

    repair_prompt = f"""The DBA asked for this tool:
{user_request}

This tool JSON failed:
{json.dumps(broken_tool, indent=2)}

Oracle returned this error:
{db_error}

Repair the tool so it stays within the allowed views and returns valid JSON with the same keys:
name, category, prompt, sql_mode, sql, queries

Output JSON only.
"""

    response = client.chat(
        oci.generative_ai_inference.models.ChatDetails(
            compartment_id=OCI_COMPARTMENT,
            chat_request=oci.generative_ai_inference.models.GenericChatRequest(
                messages=[
                    oci.generative_ai_inference.models.SystemMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=AI_TOOL_CREATOR_PROMPT)]
                    ),
                    oci.generative_ai_inference.models.UserMessage(
                        content=[oci.generative_ai_inference.models.TextContent(text=repair_prompt)]
                    ),
                ],
                max_tokens=2500,
                temperature=0.1,
                is_stream=False,
            ),
            serving_mode=oci.generative_ai_inference.models.OnDemandServingMode(
                model_id=OCI_MODEL_ID
            ),
        )
    )

    raw = response.data.chat_response.choices[0].message.content[0].text.strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("AI repair did not return valid JSON")
    return json.loads(raw[start:end + 1])

# -----------------------------
# Execution
# -----------------------------

def execute_tool(tool, time_window: dict[str, Any] | None = None, db_id: str | None = None):
    mode = tool.get("sql_mode", "single")
    collected = {}

    if mode == "single":
        sql = tool.get("sql", "")
        if not sql or not is_safe_select_sql(sql):
            raise ValueError("Tool SQL must be a safe SELECT/CTE query.")
        sql2, params = apply_time_window_sql(sql, time_window)
        collected["result"] = query_oracle(sql2, params, db_id=db_id)
    elif mode == "multi":
        queries = tool.get("queries", {})
        if not queries:
            raise ValueError("Multi-query tool has no queries defined.")
        for key, sql in queries.items():
            if not is_safe_select_sql(sql):
                raise ValueError(f"Unsafe SQL in query block: {key}")
            sql2, params = apply_time_window_sql(sql, time_window)
            collected[key] = query_oracle(sql2, params, db_id=db_id)
    else:
        raise ValueError("Unsupported tool mode")

    prompt = tool["prompt"]
    if time_window:
        prompt += f"\n\nUse this requested time window for the analysis: {time_window['label']}."

    analysis = ask_genai(prompt, collected)
    return analysis, collected

# -----------------------------
# Scenarios
# -----------------------------

SCENARIOS = [
    {
        "id": "lock",
        "name": "Lock Contention",
        "description": "Creates a blocking session automatically so you can run Blocking Sessions and Active Sessions.",
        "next_tools": ["3", "4"],
        "script": "scenario_lock_contention.py",
        "mode": "background_lock",
    },
    {
        "id": "bad_query",
        "name": "Bad Query Pattern",
        "description": "Loads AIOPS_BIG_TABLE and runs repeated full-scan style queries.",
        "next_tools": ["6", "2", "8"],
        "script": "scenario_bad_query.py",
        "mode": "run_once",
    },
    {
        "id": "cleanup",
        "name": "Cleanup / Reset",
        "description": "Resets demo tables so scenarios can be rerun cleanly.",
        "next_tools": [],
        "script": "scenario_cleanup.py",
        "mode": "run_once",
    },
]

SCENARIO_PROCS = {"lock_holder": None, "lock_waiter": None}

def _proc_running(proc):
    return proc is not None and proc.poll() is None

def get_lock_status():
    return "running" if (
        _proc_running(SCENARIO_PROCS["lock_holder"]) or
        _proc_running(SCENARIO_PROCS["lock_waiter"])
    ) else "idle"

def stop_lock_scenario():
    for key in ["lock_waiter", "lock_holder"]:
        proc = SCENARIO_PROCS.get(key)
        if _proc_running(proc):
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
        SCENARIO_PROCS[key] = None

def start_lock_scenario():
    if get_lock_status() == "running":
        return {
            "status": "running",
            "message": "Lock contention scenario is already running.",
            "next_tools": ["3", "4"]
        }

    holder_script = BASE_DIR / "scenario_lock_contention.py"
    waiter_script = BASE_DIR / "scenario_lock_waiter.py"

    if not holder_script.exists():
        raise RuntimeError("scenario_lock_contention.py not found")
    if not waiter_script.exists():
        raise RuntimeError("scenario_lock_waiter.py not found")

    holder_log = open(BASE_DIR / "lock_holder.log", "w")
    waiter_log = open(BASE_DIR / "lock_waiter.log", "w")

    holder = subprocess.Popen(
        ["python3.11", str(holder_script)],
        cwd=str(BASE_DIR),
        stdout=holder_log,
        stderr=holder_log,
        start_new_session=True,
    )
    SCENARIO_PROCS["lock_holder"] = holder
    time.sleep(3)

    if holder.poll() is not None:
        raise RuntimeError("Lock holder exited early. Check lock_holder.log")

    waiter = subprocess.Popen(
        ["python3.11", str(waiter_script)],
        cwd=str(BASE_DIR),
        stdout=waiter_log,
        stderr=waiter_log,
        start_new_session=True,
    )
    SCENARIO_PROCS["lock_waiter"] = waiter
    time.sleep(2)

    return {
        "status": "running",
        "message": "Lock contention scenario started. Go to tools 3 or 4.",
        "next_tools": ["3", "4"]
    }

# -----------------------------
# Models
# -----------------------------

class RunToolRequest(BaseModel):
    tool_id: str
    start_time: str | None = None
    end_time: str | None = None
    hours_back: int | None = None
    db_id: str | None = None

class RunScenarioRequest(BaseModel):
    scenario_id: str

class ToolUpsertRequest(BaseModel):
    id: str | None = None
    name: str
    category: str
    prompt: str
    sql_mode: str = "single"
    sql: str | None = None
    queries: dict[str, str] | None = None

class AiToolCreatorRequest(BaseModel):
    request: str

class ExportRequest(BaseModel):
    run_id: str

# -----------------------------
# Routes
# -----------------------------

@app.get("/")
def root():
    return FileResponse(STATIC_DIR / "index.html")

@app.get("/favicon.ico")
def favicon():
    return Response(status_code=204)

@app.get("/api/databases")
def api_databases():
    dbs = load_databases()
    return [{"id": d["id"], "name": d["name"], "active": d.get("active", False)} for d in dbs]

@app.get("/api/tools")
def api_tools():
    return load_tools()

@app.get("/api/tools/{tool_id}")
def api_tool_detail(tool_id: str):
    tools = load_tools()
    tool = next((t for t in tools if t["id"] == tool_id), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)
    return tool

@app.post("/api/tools/upsert")
def api_tool_upsert(req: ToolUpsertRequest):
    tools = load_tools()

    if req.category not in ["real-time", "predictive", "infra"]:
        return JSONResponse({"error": "Invalid category"}, status_code=400)

    if req.sql_mode not in ["single", "multi"]:
        return JSONResponse({"error": "SQL mode must be single or multi"}, status_code=400)

    payload = {
        "id": req.id or next_tool_id(tools),
        "name": req.name.strip(),
        "category": req.category,
        "prompt": req.prompt.strip(),
        "sql_mode": req.sql_mode
    }

    if req.sql_mode == "single":
        sql = (req.sql or "").strip()
        payload["sql"] = sql
        payload["queries"] = {}
    else:
        queries = req.queries or {}
        payload["queries"] = {k.strip(): v for k, v in queries.items() if k.strip()}
        payload["sql"] = ""

    ok, msg = validate_tool_definition(payload)
    if not ok:
        return JSONResponse({"error": msg}, status_code=400)

    payload["references"] = infer_references(
        payload["name"],
        payload["prompt"],
        payload["sql_mode"],
        payload.get("sql", ""),
        payload.get("queries", {})
    )

    existing = next((i for i, t in enumerate(tools) if t["id"] == payload["id"]), None)
    if existing is None:
        tools.append(payload)
    else:
        tools[existing] = payload

    save_tools(tools)
    return {"status": "ok", "tool": payload}

@app.delete("/api/tools/{tool_id}")
def api_tool_delete(tool_id: str):
    tools = load_tools()
    tool = next((t for t in tools if t["id"] == tool_id), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)

    tools = [t for t in tools if t["id"] != tool_id]
    save_tools(tools)
    return {"status": "ok"}

@app.post("/api/ai_tool_creator")
def api_ai_tool_creator(req: AiToolCreatorRequest):
    try:
        tool = ask_genai_json(req.request)

        draft = {
            "id": None,
            "name": (tool.get("name") or "New AI Tool").strip(),
            "category": (tool.get("category") or "real-time").strip(),
            "prompt": (tool.get("prompt") or req.request).strip(),
            "sql_mode": (tool.get("sql_mode") or "single").strip(),
            "sql": (tool.get("sql") or "").strip(),
            "queries": tool.get("queries") or {}
        }

        if draft["category"] not in ["real-time", "predictive", "infra"]:
            draft["category"] = "real-time"
        if draft["sql_mode"] not in ["single", "multi"]:
            draft["sql_mode"] = "single"

        ok, msg = validate_tool_definition(draft)
        if not ok:
            return JSONResponse({"error": msg}, status_code=400)

        # Test once against Oracle with a reasonable default window
        try:
            execute_tool(draft, build_time_window(None, None, 5))
        except Exception as first_err:
            repaired = repair_tool_with_error(req.request, draft, str(first_err))
            repaired_draft = {
                "id": None,
                "name": (repaired.get("name") or draft["name"]).strip(),
                "category": (repaired.get("category") or draft["category"]).strip(),
                "prompt": (repaired.get("prompt") or draft["prompt"]).strip(),
                "sql_mode": (repaired.get("sql_mode") or draft["sql_mode"]).strip(),
                "sql": (repaired.get("sql") or "").strip(),
                "queries": repaired.get("queries") or {}
            }
            ok2, msg2 = validate_tool_definition(repaired_draft)
            if not ok2:
                return JSONResponse({"error": f"AI repair failed validation: {msg2}"}, status_code=400)
            try:
                execute_tool(repaired_draft, build_time_window(None, None, 5))
            except Exception as second_err:
                # Second repair attempt
                repaired2 = repair_tool_with_error(req.request, repaired_draft, str(second_err))
                repaired_draft2 = {
                    "id": None,
                    "name": (repaired2.get("name") or repaired_draft["name"]).strip(),
                    "category": (repaired2.get("category") or repaired_draft["category"]).strip(),
                    "prompt": (repaired2.get("prompt") or repaired_draft["prompt"]).strip(),
                    "sql_mode": (repaired2.get("sql_mode") or repaired_draft["sql_mode"]).strip(),
                    "sql": (repaired2.get("sql") or "").strip(),
                    "queries": repaired2.get("queries") or {}
                }
                ok3, msg3 = validate_tool_definition(repaired_draft2)
                if not ok3:
                    return JSONResponse({"error": f"AI repair failed after 2 attempts: {msg3}"}, status_code=400)
                execute_tool(repaired_draft2, build_time_window(None, None, 5))
                repaired_draft = repaired_draft2
            draft = repaired_draft

        draft["references"] = infer_references(
            draft["name"],
            draft["prompt"],
            draft["sql_mode"],
            draft.get("sql", ""),
            draft.get("queries", {})
        )

        return {"status": "ok", "tool": draft}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/run")
def api_run(req: RunToolRequest):
    tools = load_tools()
    tool = next((t for t in tools if t["id"] == req.tool_id), None)
    if not tool:
        return JSONResponse({"error": "Tool not found"}, status_code=404)

    start = time.time()
    try:
        time_window = build_time_window(req.start_time, req.end_time, req.hours_back)
        analysis, raw_data = execute_tool(tool, time_window, db_id=req.db_id)
        elapsed = round(time.time() - start, 2)
        run_id = str(uuid.uuid4())

        RUN_HISTORY[run_id] = {
            "tool": tool,
            "analysis": analysis,
            "raw_data": raw_data,
            "elapsed_secs": elapsed,
            "time_window": time_window,
            "run_ts": fmt_ts(datetime.now())
        }

        return {
            "id": tool["id"],
            "name": tool["name"],
            "analysis": analysis,
            "elapsed_secs": elapsed,
            "run_id": run_id,
            "time_window": time_window
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/export/xlsx")
def api_export_xlsx(req: ExportRequest):
    run = RUN_HISTORY.get(req.run_id)
    if not run:
        return JSONResponse({"error": "Run not found"}, status_code=404)

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Tool", run["tool"]["name"]])
    ws.append(["Prompt", run["tool"]["prompt"]])
    ws.append(["Run Timestamp", run["run_ts"]])
    ws.append(["Elapsed Seconds", run["elapsed_secs"]])
    ws.append(["Time Window", run["time_window"]["label"] if run["time_window"] else "default"])
    ws.append([])
    ws.append(["AI Analysis"])
    for line in run["analysis"].splitlines():
        ws.append([line])

    for key, rows in run["raw_data"].items():
        sheet = wb.create_sheet(title=key[:31])
        if rows:
            headers = list(rows[0].keys())
            sheet.append(headers)
            for row in rows:
                sheet.append([row.get(h) for h in headers])
        else:
            sheet.append(["No rows returned"])

    out = EXPORT_DIR / f"{req.run_id}.xlsx"
    wb.save(out)
    return FileResponse(out, filename=f"{run['tool']['name'].replace(' ', '_')}.xlsx")

@app.post("/api/export/docx")
def api_export_docx(req: ExportRequest):
    run = RUN_HISTORY.get(req.run_id)
    if not run:
        return JSONResponse({"error": "Run not found"}, status_code=404)

    doc = Document()
    doc.add_heading(run["tool"]["name"], level=1)
    doc.add_paragraph(f"Prompt: {run['tool']['prompt']}")
    doc.add_paragraph(f"Run Timestamp: {run['run_ts']}")
    doc.add_paragraph(f"Elapsed Seconds: {run['elapsed_secs']}")
    doc.add_paragraph(f"Time Window: {run['time_window']['label'] if run['time_window'] else 'default'}")

    doc.add_heading("AI Analysis", level=2)
    for line in run["analysis"].splitlines():
        doc.add_paragraph(line)

    doc.add_heading("Raw Data", level=2)
    for key, rows in run["raw_data"].items():
        doc.add_heading(key, level=3)
        if not rows:
            doc.add_paragraph("No rows returned.")
            continue
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
        for row in rows[:200]:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                cells[i].text = str(row.get(h, ""))

    out = EXPORT_DIR / f"{req.run_id}.docx"
    doc.save(out)
    return FileResponse(out, filename=f"{run['tool']['name'].replace(' ', '_')}.docx")

@app.get("/api/scenarios")
def api_scenarios():
    out = []
    for s in SCENARIOS:
        status = get_lock_status() if s["id"] == "lock" else "idle"
        out.append({**s, "status": status})
    return out

@app.post("/api/run_scenario")
def api_run_scenario(req: RunScenarioRequest):
    scenario = next((s for s in SCENARIOS if s["id"] == req.scenario_id), None)
    if not scenario:
        return JSONResponse({"error": "Scenario not found"}, status_code=404)

    if scenario["id"] == "lock":
        try:
            return start_lock_scenario()
        except Exception as e:
            return JSONResponse({"error": str(e)}, status_code=500)

    script_path = BASE_DIR / scenario["script"]
    if not script_path.exists():
        return JSONResponse({"error": f"Missing script: {scenario['script']}"}, status_code=500)

    try:
        result = subprocess.run(
            ["python3.11", str(script_path)],
            cwd=str(BASE_DIR),
            capture_output=True,
            text=True,
            timeout=300,
        )
        if result.returncode != 0:
            return JSONResponse({"error": result.stderr or result.stdout or "Scenario failed"}, status_code=500)

        return {
            "status": "ok",
            "message": result.stdout.strip() or f"{scenario['name']} completed.",
            "next_tools": scenario["next_tools"]
        }
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/stop_scenario")
def api_stop_scenario(req: RunScenarioRequest):
    scenario = next((s for s in SCENARIOS if s["id"] == req.scenario_id), None)
    if not scenario:
        return JSONResponse({"error": "Scenario not found"}, status_code=404)

    try:
        if scenario["id"] == "lock":
            stop_lock_scenario()

        cleanup_script = BASE_DIR / "scenario_cleanup.py"
        if cleanup_script.exists():
            subprocess.run(
                ["python3.11", str(cleanup_script)],
                cwd=str(BASE_DIR),
                capture_output=True,
                text=True,
                timeout=120,
            )

        return {"status": "stopped", "message": f"{scenario['name']} stopped and cleanup executed."}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")